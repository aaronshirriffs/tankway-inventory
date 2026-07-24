"""
Dynamic .docx generation for the MDR Inventory API documentation.

Both the Customer Guide and the Internal Reference are generated on demand from
the live admin configuration (global settings, warehouse labels + lead times,
curated categories) so the downloads always reflect the current setup.
"""
import io
import json
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import storage
import odoo_client

API_NAME = "MDR Inventory API"
API_HOST = "api.mdrlighting.co.nz"
BODY_FONT = "Arial"
MONO_FONT = "Consolas"


def price_json(obj):
    """Pretty-print example JSON with price fields shown to 2 decimal places."""
    s = json.dumps(obj, indent=2)
    return re.sub(r'("(?:sales_price|your_price|compare_price)":\s*)(-?\d+(?:\.\d+)?)',
                  lambda m: "%s%.2f" % (m.group(1), float(m.group(2))), s)


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
def _shade(props, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def _new_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    return doc


def _heading(doc, text, size, color, before=12, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def h1(doc, t): return _heading(doc, t, 17, "12233A", before=4, after=8)
def h2(doc, t): return _heading(doc, t, 14, "1D3A8A", before=14, after=6)
def h3(doc, t): return _heading(doc, t, 12, "1B2430", before=10, after=4)


def para(doc, text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    return p


def code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line if line != "" else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _shade(p._p.get_or_add_pPr(), "F2F2F2")


def table(doc, headers, rows, header_fill="1D3A8A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        _shade(hdr[i]._tc.get_or_add_tcPr(), header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
    return t


def _save(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# live config gathering
# ---------------------------------------------------------------------------
def _source_names():
    """Map 'wh:1'/'loc:49' -> Odoo display name. Empty if Odoo unreachable."""
    names = {}
    try:
        uid, models = odoo_client.connect()
        for w in odoo_client.list_warehouses(models, uid):
            names["wh:%s" % w["id"]] = w["name"]
        for l in odoo_client.list_locations(models, uid):
            names["loc:%s" % l["id"]] = l["complete_name"]
    except Exception:
        pass
    return names


def _customer_labels():
    """Distinct customer-facing labels with their (display) lead time, from global warehouse settings."""
    wh = storage.load_warehouse_settings()
    best = {}  # label -> (rank, display_text)
    for v in wh.values():
        lbl = (v.get("label") or "").strip()
        if not lbl:
            continue
        rank = storage.lead_time_rank(v)
        if lbl not in best or rank > best[lbl][0]:
            best[lbl] = (rank, storage.lead_time_display(v))
    return sorted(((lbl, t) for lbl, (r, t) in best.items()),
                  key=lambda x: (storage.lead_time_rank(x[1]), x[0].lower()))


# ---------------------------------------------------------------------------
# Customer Guide
# ---------------------------------------------------------------------------
def build_customer_guide():
    s = storage.load_settings()
    rpm = s["default_rate_limit_per_minute"]
    rdaily = s["default_rate_limit_daily"]
    burst = s["default_burst_allowance"]
    labels = _customer_labels()
    ex = labels[:2] if labels else [("Available Immediately", ""), ("On Request", "5-7 days")]

    doc = _new_doc()
    h1(doc, "%s — Getting Started" % API_NAME)
    para(doc, "This guide explains how to connect to the %s and retrieve live "
              "stock availability using the API key we have issued to you." % API_NAME)
    para(doc, "Online version: https://%s/docs" % API_HOST, italic=True, color="444444")

    h2(doc, "Your details")
    table(doc, ["Field", "Value"], [
        ["Base URL", "https://%s" % API_HOST],
        ["Your API key", "(provided separately — keep it secret)"],
    ])

    h2(doc, "1. Authentication")
    para(doc, "Every request to /v1/products must include your API key. Use either method:")
    h3(doc, "Option A — Bearer token header (recommended)")
    code(doc, 'curl -H "Authorization: Bearer YOUR_API_KEY" \\\n     https://%s/v1/products' % API_HOST)
    h3(doc, "Option B — query parameter")
    code(doc, 'curl "https://%s/v1/products?key=YOUR_API_KEY"' % API_HOST)

    h2(doc, "2. Get products and stock — GET /v1/products")
    para(doc, "Returns the products visible to your key, each with current available "
              "stock grouped into the stock labels configured for your account.")
    h3(doc, "Which products appear")
    para(doc, "Your feed contains the products in the categories enabled for your account that are also "
              "published on the MDR website. An unpublished product will not appear, even if it is in one "
              "of your categories. Some entries are packages / kits assembled from several components — "
              "for those, the quantity shown is how many complete packages can currently be made from "
              "component stock.")
    h3(doc, "Response fields")
    table(doc, ["Field", "Type", "Description"], [
        ["id", "integer", "Stable product identifier."],
        ["name", "string", "Product name."],
        ["sku", "string (nullable)", "Product SKU / internal reference."],
        ["last_updated", "string", "UTC timestamp the product was last modified."],
        ["website_url", "string (nullable, optional)", "Link to this product's page on the MDR website, ready to use in your own listings. null if it has no public page. Only present if enabled for your key."],
        ["portal_categories", "array (optional)", "The product's B2B / portal categories from the MDR website, each as \"Parent > Name\", e.g. [\"Lighting > Moving Heads\"]. Empty array if it isn't filed under any. Only present if enabled for your key."],
        ["sales_price", "number (optional)", "Recommended retail (list) unit price. Only present if pricing is enabled for your key."],
        ["compare_price", "number (nullable, optional)", "The \"was\" / strike-through comparison price, for showing a saving against sales_price. null when the product has no comparison price. Only present if enabled for your key."],
        ["has_compare_price", "boolean (optional)", "Convenience flag — true when compare_price is set, so you can branch without null-checking."],
        ["your_price", "number (optional)", "Your account's unit buy price (qty 1) under your pricelist. Only present if a pricelist is assigned to your key."],
        ["stock", "object", "Available quantity by label, e.g. {\"Available Immediately\": 12}. Free to sell — stock already reserved against confirmed orders is excluded."],
        ["availability", "array", "Same quantities per label, each with a lead time: {label, qty, lead_time}."],
        ["incoming", "array (optional)", "Stock on a scheduled inbound delivery not yet arrived: {quantity, expected_date, label}, where expected_date is the delivery's scheduled arrival date. Only present if enabled for your key."],
    ])

    h3(doc, "Example response")
    sample_stock = {lbl: q for (lbl, _), q in zip(ex, [7, 3])}
    sample_avail = [{"label": lbl, "qty": q, "lead_time": lead}
                    for ((lbl, lead), q) in zip(ex, [7, 3])]
    example = {
        "count": 1,
        "products": [{
            "id": 1423,
            "name": "Chauvet Maverick MK3 Spot",
            "sku": "CHV-MK3-SPOT",
            "last_updated": "2026-06-05 21:14:02",
            "website_url": "https://www.mdrlighting.co.nz/shop/chauvet-maverick-mk3-spot-1423",
            "portal_categories": ["Lighting > Moving Heads", "Brands > Chauvet"],
            "sales_price": 4299.00,
            "compare_price": 4799.00,
            "has_compare_price": True,
            "your_price": 3869.10,
            "stock": sample_stock,
            "availability": sample_avail,
            "incoming": [{"quantity": 20, "expected_date": "2026-06-15",
                          "label": (ex[0][0] if ex else "Incoming")}],
        }],
    }
    code(doc, price_json(example))

    h3(doc, "Availability & lead times")
    para(doc, "The stock object and the availability array describe the same quantities. "
              "availability adds a lead_time per label — a short text such as \"5-7 days\"; an empty "
              "value means the item is available now and ready to ship.")
    para(doc, "These quantities are available (free-to-sell) stock, not raw shelf count: anything already "
              "reserved against a confirmed order has been deducted, so the number you see is what can "
              "actually be ordered today.")
    if labels:
        para(doc, "Stock labels currently configured:")
        table(doc, ["Stock label", "Lead time"],
              [[lbl, (lead or "In stock now")] for lbl, lead in labels])

    h3(doc, "Incoming stock")
    para(doc, "If incoming stock is enabled for your key, each product includes an incoming array of "
              "inbound deliveries that are scheduled but not yet received — each with a quantity, an "
              "expected_date (YYYY-MM-DD, the delivery's scheduled arrival date) and a destination label. "
              "A product with nothing due in returns an empty array. Use this to tell customers when "
              "out-of-stock items are due back in.")
    para(doc, "Only live, scheduled deliveries are listed: a cancelled delivery disappears from incoming, "
              "and a re-planned arrival date is reflected here — so dates track the current delivery "
              "schedule rather than the date originally ordered.")

    h2(doc, "3. Health check — GET /v1/status (no key required)")
    code(doc, json.dumps({"api": API_NAME, "status": "ok",
                          "timestamp": "2026-06-06T14:21:09+12:00"}, indent=2))

    h2(doc, "4. Error codes")
    table(doc, ["Status", "Meaning"], [
        ["401", "Missing, invalid, disabled, or expired API key."],
        ["429", "Rate limit exceeded — retry after retry_after_seconds."],
        ["500", "Unexpected server error."],
    ])

    h2(doc, "5. Rate limits")
    para(doc, "Your key has a per-minute limit, a short burst allowance, and a daily cap. "
              "Typical limits are around %d requests/minute, a burst of %d, and %d per day "
              "(your key may differ). On a 429, wait retry_after_seconds and retry."
              % (rpm, burst, rdaily))
    code(doc,
         "import time, requests\n"
         "URL = \"https://%s/v1/products\"\n"
         "HEADERS = {\"Authorization\": \"Bearer YOUR_API_KEY\"}\n\n"
         "def get_products():\n"
         "    while True:\n"
         "        r = requests.get(URL, headers=HEADERS, timeout=30)\n"
         "        if r.status_code == 429:\n"
         "            time.sleep(r.json().get(\"retry_after_seconds\", 30)); continue\n"
         "        r.raise_for_status(); return r.json()" % API_HOST)

    h2(doc, "6. Tips")
    bullet(doc, "Cache results for a few minutes rather than calling on every page load.")
    bullet(doc, "Match on sku (or id) when syncing to your own catalogue; id is the most stable identifier.")
    bullet(doc, "Show lead times to your buyers using availability[].lead_time.")

    para(doc, "For key requests or support, contact your MDR Lighting account manager.",
         italic=True, color="444444")
    return _save(doc)


# ---------------------------------------------------------------------------
# Internal Reference
# ---------------------------------------------------------------------------
def build_internal_reference():
    s = storage.load_settings()
    rpm = s["default_rate_limit_per_minute"]
    rdaily = s["default_rate_limit_daily"]
    burst = s["default_burst_allowance"]
    wh = storage.load_warehouse_settings()
    names = _source_names()
    available = storage.load_available_categories()

    doc = _new_doc()
    h1(doc, "%s — Internal Reference Guide" % API_NAME)
    para(doc, "Generated live from the current admin configuration. Lead times, labels, "
              "rate-limit defaults and category presets below reflect the system as configured now.")

    h2(doc, "1. Surfaces")
    table(doc, ["Surface", "URL", "Auth"], [
        ["Admin panel", "tools.tankway.co.nz/inventory/", "Login (admins)"],
        ["Public API", "https://%s/v1/..." % API_HOST, "API key"],
    ])

    h2(doc, "2. Per-key controls")
    table(doc, ["Setting", "What it does"], [
        ["Allowed categories", "Which Odoo categories are visible (sub-categories included)."],
        ["Excluded categories", "Hide a sub-category subtree even within an allowed parent."],
        ["Excluded SKUs", "Hide specific product codes."],
        ["Warehouse / location mappings", "Group warehouses and EXT partner locations under customer labels."],
        ["Pricelist", "Assign an Odoo pricelist; adds your_price (customer buy price, qty 1) per product."],
        ["Show price", "Include sales_price (list price) in the response."],
        ["Show compare price", "Include compare_price + has_compare_price (the \"was\"/strike-through price). Off by default."],
        ["Show website URL", "Include website_url, the product's public page on the MDR website. Off by default."],
        ["Show B2B portal categories", "Include portal_categories — the website portal categories as [\"Parent > Name\"]. Same data as the export's \"B2B portal category/s\" column. Off by default."],
        ["Show incoming stock", "Include the incoming array (scheduled inbound deliveries not yet received; cancelled ones excluded). Off by default."],
        ["Expiry", "Optional date after which the key stops working."],
        ["Rate limits", "Per-minute, burst, daily caps (defaults below)."],
    ])

    h2(doc, "3. Default rate limits for new keys")
    para(doc, "Configured on the Settings page. Current values:")
    table(doc, ["Limit", "Value"], [
        ["Per minute", rpm],
        ["Daily cap", rdaily],
        ["Burst allowance", burst],
    ])

    h2(doc, "4. Warehouses & lead times")
    para(doc, "Global label and delivery lead time per stock source. Lead times surface in the "
              "API availability field.")
    wh_rows = []
    if wh:
        for key, v in sorted(wh.items()):
            wh_rows.append([
                names.get(key, key),
                v.get("label", "") or "—",
                storage.lead_time_display(v) or "—",
            ])
    if wh_rows:
        table(doc, ["Stock source", "Default label", "Lead time"], wh_rows)
    else:
        para(doc, "No warehouse labels/lead times configured yet (Warehouses page).", italic=True, color="666666")

    h2(doc, "5. Available categories")
    if available:
        para(doc, "Key-generation pickers are limited to these curated categories:")
        table(doc, ["Category"], [[c.get("complete_name", str(c.get("id")))] for c in available])
    else:
        para(doc, "No curation set — all Odoo categories are available in the key-generation pickers.",
             italic=True, color="666666")

    h2(doc, "6. Public API response shape")
    code(doc, price_json({
        "count": 1,
        "products": [{
            "id": 1234, "name": "Maestro DMX Controller", "sku": "0011767",
            "last_updated": "2026-06-01 09:12:33",
            "stock": {"Auckland": 12.0, "AUS Partner": 3.0},
            "availability": [
                {"label": "Auckland", "qty": 12.0, "lead_time": ""},
                {"label": "AUS Partner", "qty": 3.0, "lead_time": "5-7 days"},
            ],
            "incoming": [
                {"quantity": 20, "expected_date": "2026-06-15", "label": "Auckland"},
            ],
            "website_url": "https://www.mdrlighting.co.nz/shop/example-product-1423",
            "sales_price": 499.0,
            "compare_price": 599.0,
            "has_compare_price": True,
            "your_price": 449.10,
        }],
    }))
    para(doc, "incoming is present only when 'Show incoming stock' is enabled for the key "
              "(empty array when nothing is due in); label maps the destination warehouse to the "
              "key's customer-facing label, falling back to \"Incoming\". It is built from scheduled "
              "inbound deliveries, so a cancelled delivery drops out and a re-planned date is reflected. "
              "website_url / compare_price + has_compare_price appear only when their toggles are on; "
              "stock is available (free-to-sell) quantity, excluding stock reserved against confirmed orders.",
              italic=True, color="666666")

    h2(doc, "7. Admin sidebar pages")
    table(doc, ["Page", "Purpose"], [
        ["API Keys", "Manage per-customer keys."],
        ["Activity Log", "Per-customer API usage summary; click a customer for recent requests. Persisted across restarts."],
        ["Status", "Service health, Odoo connectivity, counts."],
        ["Warehouses", "Global labels + lead times per source."],
        ["Categories", "Curate available categories."],
        ["Settings", "Default rate limits for new keys; documentation downloads."],
    ])

    h2(doc, "8. Files & operations")
    table(doc, ["File", "Role"], [
        ["app.py / odoo_client.py / storage.py / rate_limit.py", "Application code."],
        ["docgen.py", "Generates these documents from live config."],
        ["keys.json", "API keys (source of truth)."],
        ["warehouse_settings.json / category_settings.json / settings.json", "Global config."],
        [".env", "Secrets (Odoo + Flask)."],
    ])
    code(doc, "systemctl restart inventory.service\n"
              "journalctl -u inventory.service -f\n"
              "curl https://%s/v1/status" % API_HOST)
    return _save(doc)
